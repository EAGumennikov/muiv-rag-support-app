from __future__ import annotations

"""
Сервис формирования экспортных документов.

Модуль отвечает за прикладные выгрузки `.docx` и `.xlsx` для web-сервиса:
ответы RAG, статьи базы знаний, пользовательская история и административные
отчеты. Файлы формируются в памяти через BytesIO, чтобы не оставлять
runtime-документы в репозитории и не смешивать пользовательские выгрузки с кодом.
"""

import re
import unicodedata
from datetime import datetime
from io import BytesIO
from typing import Any, Iterable

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def build_rag_answer_docx(answer: dict[str, Any]) -> tuple[BytesIO, str]:
    # Документ с RAG-ответом предназначен для передачи результата работы
    # поддержки в удобном виде: вопрос, ответ и понятные пользователю источники.
    document = Document()
    document.add_heading("RAG-ответ службы поддержки", level=1)
    _add_key_value(document, "Дата формирования", _current_timestamp())
    _add_key_value(document, "Дата исходного запроса", answer.get("query_created_at") or "Не указана")
    _add_key_value(document, "Пользователь", answer.get("username") or "Не указан")

    document.add_heading("Вопрос", level=2)
    _add_markdown_text(document, answer.get("question_text", ""))

    document.add_heading("Ответ", level=2)
    _add_markdown_text(document, answer.get("answer_text", ""))

    document.add_heading("Источники", level=2)
    sources = answer.get("sources") or []
    if not sources:
        document.add_paragraph("Источники не зафиксированы.")
    for index, source in enumerate(sources, start=1):
        document.add_paragraph(f"{index}. {source.get('title') or 'Источник базы знаний'}", style="List Number")
        if source.get("breadcrumbs"):
            _add_key_value(document, "Размещение", " > ".join(source["breadcrumbs"]))
        if source.get("section_label"):
            _add_key_value(document, "Раздел", source["section_label"])
        if source.get("original_url"):
            _add_key_value(document, "Оригинал", source["original_url"])
        if source.get("doc_id"):
            _add_key_value(document, "Внутренняя ссылка", source["doc_id"])

    return _save_docx(document), _safe_filename(f"rag-answer-{answer.get('id', 'export')}", ".docx")


def build_article_docx(article: dict[str, Any]) -> tuple[BytesIO, str]:
    # Экспорт статьи публичен по той же логике, что и страница /article:
    # в документ попадают только пользовательские поля, без source_file и путей.
    document = Document()
    document.add_heading(article.get("title") or "Статья базы знаний", level=1)
    _add_key_value(document, "Дата формирования", _current_timestamp())
    if article.get("breadcrumbs"):
        _add_key_value(document, "Размещение", " > ".join(article["breadcrumbs"]))
    if article.get("author"):
        _add_key_value(document, "Автор", article["author"])
    if article.get("date"):
        _add_key_value(document, "Дата статьи", article["date"])
    if article.get("original_url"):
        _add_key_value(document, "Оригинал", article["original_url"])

    document.add_heading("Материал", level=2)
    _add_markdown_text(document, article.get("body_markdown_text") or article.get("markdown_text") or "")

    return _save_docx(document), _safe_filename(article.get("title") or article.get("doc_id") or "article", ".docx")


def build_user_history_xlsx(rows: list[dict[str, Any]], username: str) -> tuple[BytesIO, str]:
    # Персональная история включает только записи текущего пользователя:
    # дата, вопрос, метрики retrieval и краткий фрагмент ответа.
    headers = ["Дата запроса", "Вопрос", "Найдено источников", "Использовано источников", "Ответ есть", "Фрагмент ответа"]
    values = [
        [
            row.get("created_at", ""),
            row.get("question_text", ""),
            row.get("retrieved_chunks_count", 0),
            row.get("used_chunks_count", 0),
            row.get("has_answer", ""),
            row.get("answer_excerpt", ""),
        ]
        for row in rows
    ]
    return _single_sheet_workbook("Моя история запросов", headers, values), _safe_filename(f"history-{username}", ".xlsx")


def build_feedback_xlsx(rows: list[dict[str, Any]]) -> tuple[BytesIO, str]:
    headers = ["Дата", "Пользователь", "Имя", "Email", "Тема", "Текст обращения", "Краткий фрагмент", "Статус"]
    values = [
        [
            row.get("created_at", ""),
            row.get("username", "") or "Не авторизован",
            row.get("name", ""),
            row.get("email", ""),
            row.get("topic", ""),
            row.get("message", ""),
            row.get("message_excerpt", ""),
            row.get("status", ""),
        ]
        for row in rows
    ]
    return _single_sheet_workbook("Feedback", headers, values), "feedback.xlsx"


def build_search_history_xlsx(rows: list[dict[str, Any]]) -> tuple[BytesIO, str]:
    headers = ["Дата", "Пользователь", "Вопрос", "Найдено источников", "Использовано источников", "Ответ есть", "Фрагмент ответа"]
    values = [
        [
            row.get("created_at", ""),
            row.get("username", "") or "Не авторизован",
            row.get("question_text", ""),
            row.get("retrieved_chunks_count", 0),
            row.get("used_chunks_count", 0),
            row.get("has_answer", ""),
            row.get("answer_excerpt", ""),
        ]
        for row in rows
    ]
    return _single_sheet_workbook("История запросов", headers, values), "search-history.xlsx"


def build_admin_statistics_xlsx(payload: dict[str, Any]) -> tuple[BytesIO, str]:
    # Статистическая книга содержит несколько листов: сводные счетчики,
    # распределение feedback по статусам и последние события аудита.
    workbook = Workbook()
    summary_sheet = workbook.active
    summary_sheet.title = "Сводка"
    _fill_sheet(summary_sheet, ["Показатель", "Значение"], [[row["metric"], row["value"]] for row in payload.get("summary", [])])

    status_sheet = workbook.create_sheet("Статусы feedback")
    _fill_sheet(status_sheet, ["Статус", "Количество"], [[row["status"], row["count"]] for row in payload.get("feedback_statuses", [])])

    audit_sheet = workbook.create_sheet("Аудит")
    _fill_sheet(
        audit_sheet,
        ["Дата", "Событие", "Сущность", "ID"],
        [
            [row.get("created_at", ""), row.get("event_type", ""), row.get("entity_type", ""), row.get("entity_id", "")]
            for row in payload.get("audit", [])
        ],
    )
    return _save_workbook(workbook), "admin-statistics.xlsx"


def _single_sheet_workbook(title: str, headers: list[str], rows: list[list[Any]]) -> BytesIO:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = title[:31]
    _fill_sheet(sheet, headers, rows)
    return _save_workbook(workbook)


def _fill_sheet(sheet, headers: list[str], rows: list[list[Any]]) -> None:
    # Таблица оформляется минимально: жирная шапка, перенос строк и ширина
    # колонок, чтобы выгрузка была читаемой сразу после скачивания.
    header_fill = PatternFill(fill_type="solid", fgColor="E7D6BF")
    for column_index, header in enumerate(headers, start=1):
        cell = sheet.cell(row=1, column=column_index, value=header)
        cell.font = Font(bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(wrap_text=True, vertical="top")

    for row_index, row in enumerate(rows, start=2):
        for column_index, value in enumerate(row, start=1):
            cell = sheet.cell(row=row_index, column=column_index, value=value)
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    sheet.freeze_panes = "A2"
    for column_index, header in enumerate(headers, start=1):
        values = [str(header)] + [str(row[column_index - 1]) if column_index - 1 < len(row) else "" for row in rows]
        width = min(max(len(value) for value in values) + 2, 64)
        sheet.column_dimensions[get_column_letter(column_index)].width = max(width, 12)


def _add_key_value(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(str(value or "Не указано"))


def _add_markdown_text(document: Document, markdown_text: str) -> None:
    # Для docx используем легкую markdown-нормализацию без HTML:
    # заголовки, списки и обычные абзацы. Служебная разметка не выводится.
    lines = (markdown_text or "").replace("\r\n", "\n").splitlines()
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        heading_level = _heading_level(stripped)
        if heading_level:
            document.add_heading(stripped.lstrip("#").strip(), level=min(heading_level + 1, 4))
            continue
        if stripped.startswith(("- ", "* ")):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue
        if re.match(r"^\d+[.)]\s+", stripped):
            document.add_paragraph(re.sub(r"^\d+[.)]\s+", "", stripped), style="List Number")
            continue
        document.add_paragraph(stripped)


def _heading_level(line: str) -> int:
    match = re.match(r"^(#{1,6})\s+", line)
    return len(match.group(1)) if match else 0


def _save_docx(document: Document) -> BytesIO:
    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


def _save_workbook(workbook: Workbook) -> BytesIO:
    output = BytesIO()
    workbook.save(output)
    output.seek(0)
    return output


def _current_timestamp() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _safe_filename(value: str, extension: str) -> str:
    # Имя файла очищается для HTTP-заголовка и файловых систем. Кириллицу
    # транслитерировать не пытаемся: оставляем ASCII-дублер, как в markdown.
    normalized = unicodedata.normalize("NFKD", value or "export").encode("ascii", "ignore").decode("ascii")
    normalized = re.sub(r"[^A-Za-z0-9._-]+", "_", normalized).strip("._")
    if not normalized:
        normalized = "export"
    if not normalized.lower().endswith(extension):
        normalized += extension
    return normalized
